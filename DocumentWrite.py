import docx
from docx import Document  # Document module from docx library
from docx.shared import Pt # import Pt module from docx.shared module, this is for font sizes in points
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Specify text alignment for paragraphs
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml # It's used to parse and generate custom XML elements for advanced document manipulation.
from docx.shared import RGBColor # RGBColor is used to specify text and background colors in a document.

# class for writing the data to the word document
class DocumentWrite:
  def __init__(self, user_cart_dict, query_cart_item_data, customer_user_id, customer_cart_id):
    self.doc = Document()                            # create document object
    self.user_cart_dict = user_cart_dict             # customer details dictionary
    self.query_cart_item_data = query_cart_item_data # customer cart items list of tuples
    self.user_id = customer_user_id                  # customer user id
    self.cart_id = customer_cart_id                  # customer cart id
      
  def write_quote_number(self):
    # check whether the length of the user_id is less than six or not
    if len(self.user_id) < 6:
    # Add zeros at the beginning of user_id until it becomes 6 characters long
      self.user_id = self.user_id.zfill(6)  
    # check whether the length of the cart_id is less than six or not
    if len(self.cart_id) < 6:
      # Add zeros at the beginning of cart_id until it becomes 6 characters long
      self.cart_id = self.cart_id.zfill(6)

  # function for word document operations
  def open_document(self):
    self.doc = Document()                           # create a document object
    title = self.doc.add_heading(f'Invoice for Quote: {self.user_id}{self.cart_id}', level=0) # heading combination of user id and cart id
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # align text to the center
    title.runs[0].bold = True                       # make text as bold
    title.runs[0].font.size = Pt(14)                # set font size to 14

  # customer information table
  def write_data_to_first_table(self):
    # Keys we need for printing the informations on table
    self.dict_keys = ['CUSTOMER EMAIL-ADDR', 'CUSTOMER COMPANY', 'CUSTOMER ADDR1', 
                      'CUSTOMER ADDR2','CUSTOMER CITY','CUSTOMER STATE', 'CUSTOMER PROVINCE',
                      'CUSTOMER COUNTRY','CUSTOMER ZIP-CODE','CUSTOMER BUSINESS PHONE','CUSTOMER CODE']
    
    # Values we are interested
    self.dict_values = [self.user_cart_dict[key] for key in self.dict_keys]

    # Keys
    self.table_details = ['To:', 'Company:', 'Address Line 1','Address Line 2','City','State', 'Province',
                          'Country','Postal code','Business phone','Customer system code']
    
    self.table = self.doc.add_table(rows = len(self.table_details), cols = 2, style = 'Table Grid') # create a table with 2 columns and total number of rows are equals to length of the keys
    
    for i in range(len(self.table_details)):
      self.table.cell(i, 0).text = self.table_details[i] # first columns, print the keys
      self.table.cell(i, 1).text = self.dict_values[i]   # second columns, print the values
    

    # Set the entire first column to bold
    for row in self.table.rows:     # iterates through each row in the table.
      #For each row in the table, we retrieve the first cell (column 0) in that row.
      self.first_cell = row.cells[0]  
      # We then iterate through each paragraph within the cell. A paragraph is a block of text,and cells can contain multiple paragraphs,
      for paragraph in self.first_cell.paragraphs:
        #  For each paragraph in the cell, we iterate through the runs.
        #  A run is a continuous sequence of characters with the same formatting. 
        for run in paragraph.runs: 
          # Within each run, we set the bold property to True. 
          # This makes the text within the run (the specific portion of the paragraph) bold
          run.bold = True
  
  # function for adding the paragraph after the table
  def add_first_paragraph(self):
    self.para = self.doc.add_paragraph()
    run = self.para.add_run()
    run.add_break()
    self.para.paragraph_format.line_spacing = 1.5
    self.user_firstname = self.user_cart_dict.get('CUSTOMER FIRSTNAME')
    self.user_lastname = self.user_cart_dict.get('CUSTOMER LASTNAME')
    static_text = f'''Dear Mr. {self.user_firstname }{self.user_lastname}
Thank you for purchasing goods using our services,
Please confirm that the ordered items are correct and proceed with payment using one of our 
payment services as per agreement
    '''
    self.para.add_run(static_text).italic = True
    self.doc.add_page_break()


  # function for creating a order table 
  def order_table_create(self):
    self.para = self.doc.add_paragraph()
    run = self.para.add_run(' You have ordered the following items:')
    run.underline = True
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Cambria'
    run.font.color.rgb = RGBColor(128, 128, 128)
    self.doc.add_paragraph()
    # first row elements ==> headings
    self.header = ['Catalog Code', 'Description\n\nUser Description\n\nDescription Long', 'Quantity',
                    'Net price\n\nRecurring Price','Currency', 'Product Weight', 'Unit of Measure',
                      'Item Delivery Status\n\nItem Delivery Method']
    self.query_items = []
    for item in self.query_cart_item_data:
      self.query_cart = [item[0],"\n\n".join(item[1:4]), item[4], 
                                 "\n\n".join(map(str, item[5:7])), item[7], item[8], item[9],
                                 "\n\n".join(map(str, item[10:]))]
      self.query_items.append(self.query_cart)
    print(self.query_items)
    # Add a table to the document
    self.num_rows = len(self.query_items) + 1 
    self.num_cols = 8
    self.table = self.doc.add_table(rows = self.num_rows, cols = self.num_cols, style = 'Table Grid')

    if self.num_rows == 1:  # Check if there are items in the cart
      for i in range( self.num_cols):
        self.table.cell(0, i).text = self.header[i]  # adding heding to the first rows
    else:
      for i in range(self.num_cols):
        self.table.cell(0, i).text = self.header[i] # adding heding to the first rows
      for i in range(1, self.num_rows):
        for j in range(self.num_cols):
          self.table.cell(i, j).text = str(self.query_items[i - 1][j])
         
    # make first row (title) elements bold
    for cell in self.table.rows[0].cells:
      for paragraph in cell.paragraphs:
        for run in paragraph.runs:
          run.bold= True

    self.doc.add_page_break()
  # function for adding the seconf paragraph
  def add_second_paragraph(self):
    if self.query_cart_item_data:
      self.doc.add_paragraph() 
      self.para = self.doc.add_paragraph() 
      self.para.add_run(f"A special quote has been created as per agreement for you on: ")
      run = self.para.add_run(f'{self.user_cart_dict.get("DATE ORDERED")}')
      run.bold = True
      run.font.color.rgb = RGBColor(0, 150, 255)

      self.para = self.doc.add_paragraph() 
      self.para.add_run("Your quote currently is in status: ")
      run = self.para.add_run(f'{self.user_cart_dict.get("ORDER STATUS")}')
      run.bold = True
      run.font.color.rgb = RGBColor(0, 150, 255)
      self.para.add_run(" ordered Items will be delivered by: ")
      run = self.para.add_run(f'{self.user_cart_dict.get("SHIPPING METHOD")}')
      run.bold = True
      run.font.color.rgb = RGBColor(0, 150, 255)

      self.para = self.doc.add_paragraph() 
      self.para.add_run("There are some comments made by our sales manager:  ")
      run = self.para.add_run(f'{self.user_cart_dict.get("CART COMMENT")}')
      run.bold = True
      run.font.color.rgb = RGBColor(0, 150, 255)
    else:
      self.doc.add_paragraph('Nothing in the cart') 


  # function for creating third table 
  def write_data_to_third_table(self):
    # Keys we need for printing the information on the table
    self.dict_keys = ['SHIPPING_COST', 'PAYMENT METHOD', 'PAYMENT APROOVED', 'TOTAL AMOUNT', 'CURRENCY QUOTE', 'SEND EMAIL END USER', 'ORIGIN',
                      'MY SHIP TO', 'MY BILL TO']

    # Values we are interested in
    self.dict_values = [str(self.user_cart_dict[key]) for key in self.dict_keys]

    # Keys
    self.table_details = ['Shipping Cost:', 'Payment method:', 'Payment approved', 'Total amount', 'Currency quoted in', 'Has email been sent', 'Quote generated from',
                          'Ship to address', 'Bill to address']

    self.table = self.doc.add_table(rows=len(self.table_details), cols=2, style = 'Table Grid')

    for i in range(len(self.table_details)):
        self.table.cell(i, 0).text = self.table_details[i]
        self.table.cell(i, 1).text = self.dict_values[i]

    # Set the entire first column to bold
    for row in self.table.rows:     
      self.first_cell = row.cells[0]  
      for paragraph in self.first_cell.paragraphs:
        for run in paragraph.runs: 
          run.bold = True
  
  # function for saving the word document
  def save_document(self):
    '''
    To add text or images to a header or footer in python-docx, 
    add a paragraph to the header or footer 
    and then add a run to that paragraph.
    '''
    self.section = self.doc.sections[0]        # Access the first section of the document
    self.footer = self.section.footer          # add a footer to the section
    self.paragraph = self.footer.paragraphs[0] # add a paragraph to the footer, paragraphs[0] index corresponds to an existing or new paragraph in the footer section. 
    self.run = self.paragraph.add_run()        # add a run to the paragraph
    self.run.add_picture('C:\\Users\\admin\\Desktop\\BOOTCAMP\\clarity-log.png', width = docx.shared.Inches(1.25)) # insert the clarity log, specifying the size 
    self.doc.save(f"Outout_doc_for_{self.user_id}{self.cart_id}.docx")
